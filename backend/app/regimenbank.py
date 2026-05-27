from __future__ import annotations

import calendar as cal
import datetime as dt
import re
from dataclasses import dataclass, asdict, field, replace
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple


@dataclass
class TherapyOption:
    dose: str
    duration: str
    total_doses: Optional[int] = None


@dataclass
class Chemotherapy:
    name: str
    route: str
    dose: str
    frequency: str
    duration: str
    total_doses: Optional[int] = None
    options: List[TherapyOption] = field(default_factory=list)

    @staticmethod
    def from_dict(d: Dict[str, Any]) -> "Chemotherapy":
        opts_data = d.get("options", [])
        parsed_opts = [
            TherapyOption(
                dose=o.get("dose", ""),
                duration=o.get("duration", ""),
                total_doses=o.get("total_doses"),
            )
            for o in opts_data
        ]
        return Chemotherapy(
            d["name"], d["route"], d["dose"], d["frequency"], d["duration"], d.get("total_doses"), parsed_opts
        )


@dataclass
class Regimen:
    name: str
    disease_state: Optional[str] = None
    on_study: bool = False
    notes: Optional[str] = None
    therapies: List[Chemotherapy] = field(default_factory=list)

    @staticmethod
    def from_dict(name: str, d: Dict[str, Any]) -> "Regimen":
        return Regimen(
            name=name, disease_state=d.get("disease_state"), on_study=d.get("on_study", False),
            notes=d.get("notes"), therapies=[Chemotherapy.from_dict(x) for x in d.get("therapies", [])],
        )

    def to_dict(self) -> Dict[str, Any]:
        return {
            "disease_state": self.disease_state, "on_study": self.on_study,
            "notes": self.notes, "therapies": [asdict(t) for t in self.therapies],
        }

    def upsert_chemo(self, c: Chemotherapy) -> None:
        key = c.name.strip().lower()
        for i, t in enumerate(self.therapies):
            if t.name.strip().lower() == key:
                self.therapies[i] = c
                return
        self.therapies.append(c)


def parse_day_spec(day_spec: str) -> List[int]:
    if not day_spec:
        return []
    s = day_spec.replace("–", "-").strip().lower()
    s = re.sub(r"^days?\s*[:\-]?\s*", "", s)
    if not s:
        return []

    tokens = re.split(r"[,\s]+", s)
    out: List[int] = []

    for tok in tokens:
        if not tok:
            continue
        if "-" in tok:
            try:
                a_str, b_str = tok.split("-", 1)
                a, b = int(a_str), int(b_str)
                if a <= b:
                    out.extend(range(a, b + 1))
            except ValueError:
                continue
        else:
            try:
                out.append(int(tok))
            except ValueError:
                continue

    return sorted(set(d for d in out if d >= 1))


def _doses_per_day(frequency: str) -> int:
    """Return the number of doses given per active day based on frequency abbreviation."""
    if not frequency:
        return 1
    f = frequency.strip().lower()
    # Four times daily
    if re.search(r"\bqid\b|q\.?6\.?h\b|four\s+times\s+daily|4\s*x\s*(daily|day)", f):
        return 4
    # Three times daily
    if re.search(r"\btid\b|q\.?8\.?h\b|three\s+times\s+daily|3\s*x\s*(daily|day)", f):
        return 3
    # Twice daily
    if re.search(r"\bbid\b|q\.?12\.?h\b|twice\s+daily|twice\s+a\s+day|2\s*x\s*(daily|day)", f):
        return 2
    return 1


def compute_calendar_grid(reg: Regimen, start: dt.date, cycle_len: int):
    max_day = cycle_len
    by_day: Dict[int, List[str]] = {d: [] for d in range(1, cycle_len + 1)}
    for t in reg.therapies:
        dlist = [d for d in parse_day_spec(t.duration) if d <= cycle_len]
        if dlist:
            max_day = max(max_day, max(dlist))
            for d in dlist:
                by_day.setdefault(d, []).append(t.name)

    first_sun = start - dt.timedelta(days=(start.weekday() + 1) % 7)
    last_needed = start + dt.timedelta(days=max_day - 1)
    last_sat = last_needed + dt.timedelta(days=(5 - last_needed.weekday()) % 7)

    grid: List[List[Dict[str, Any]]] = []
    d = first_sun
    week: List[Dict[str, Any]] = []
    while d <= last_sat:
        entry: Dict[str, Any] = {"date": d, "cycle_day": None, "labels": []}
        if d >= start:
            cd = (d - start).days + 1
            if 1 <= cd <= max_day:
                entry["cycle_day"] = cd
                entry["labels"] = by_day.get(cd, []) or ["Rest"]
        week.append(entry)
        if len(week) == 7:
            grid.append(week)
            week = []
        d += dt.timedelta(days=1)
    if week:
        while len(week) < 7:
            week.append({"date": d, "cycle_day": None, "labels": []})
            d += dt.timedelta(days=1)
        grid.append(week)
    return first_sun, last_sat, max_day, grid


def _spell_route(route: str) -> str:
    r = route.strip().upper()
    mapping = {"PO": "by mouth", "IV": "intravenously", "SQ": "Inject SubQ", "IT": "Given during lumbar puncture"}
    return mapping.get(r, route)


def export_calendar_docx(reg: Regimen, start: dt.date, cycle_len: int, out_path: Path, cycle_label: str, note: Optional[str] = None) -> bool:
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return False

    first_sun, last_sat, _, grid = compute_calendar_grid(reg, start, cycle_len)
    months = cal.month_name[first_sun.month]
    if first_sun.month != last_sat.month or first_sun.year != last_sat.year:
        months += f" - {cal.month_name[last_sat.month]}"
    year = str(first_sun.year) if first_sun.year == last_sat.year else f"{first_sun.year}-{last_sat.year}"

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    sec = doc.sections[0]
    sec.orientation = 1
    sec.page_width, sec.page_height = Inches(11), Inches(8.5)
    sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Inches(0.5)

    htbl = doc.add_table(rows=1, cols=2)
    htbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    left_cell = htbl.cell(0, 0)
    left_cell.text = ""
    p_name = left_cell.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(0)
    r_name = p_name.add_run("First Last")
    r_name.italic = True
    r_name.font.size = Pt(12)

    p_dob = left_cell.add_paragraph()
    p_dob.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_dob.paragraph_format.space_before = Pt(0)
    p_dob.paragraph_format.space_after = Pt(0)
    r_dob = p_dob.add_run("DOB: M/DD/YYYY")
    r_dob.italic = True
    r_dob.font.size = Pt(12)

    if left_cell.paragraphs and left_cell.paragraphs[0].text == "":
        try:
            left_cell._element.remove(left_cell.paragraphs[0]._element)
        except Exception:
            pass

    right_cell = htbl.cell(0, 1)
    right_p = right_cell.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    logo_path = None
    possible_paths = [
        Path(__file__).resolve().parent.parent / "ucm.png",
        Path(__file__).resolve().parent / "ucm.png",
        Path.cwd() / "ucm.png",
        Path.cwd() / "backend" / "ucm.png",
        Path("/app/ucm.png"),
    ]
    for p in possible_paths:
        if p.exists():
            logo_path = p
            break

    if logo_path:
        try:
            right_p.add_run().add_picture(str(logo_path), height=Inches(0.76))
        except Exception as e:
            print(f"[WARN] Could not insert logo: {e}")

    doc.add_paragraph()

    rows = len(grid) + 2
    cols = 7
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    title_row = table.rows[0]
    title_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    title_row.height = Inches(0.72)

    first_cell = title_row.cells[0]
    for j in range(1, cols):
        first_cell.merge(title_row.cells[j])
    c = first_cell
    c.text = ""

    p1 = c.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run("Chemotherapy Calendar")
    r1.bold = True
    r1.font.size = Pt(14)

    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)

    r2a = p2.add_run(reg.name + " - ")
    r2a.font.size = Pt(14)

    r2b = p2.add_run(cycle_label)
    r2b.bold = True
    r2b.font.size = Pt(14)

    p3 = c.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(0)
    r3 = p3.add_run(f"{months} {year}")
    r3.font.size = Pt(14)

    if note:
        p4 = c.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.paragraph_format.space_before = Pt(0)
        p4.paragraph_format.space_after = Pt(0)
        r4 = p4.add_run(note)
        r4.italic = True
        r4.font.size = Pt(11)

    header_names = ["Sunday", "Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday"]
    header_row = table.rows[1]
    header_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    header_row.height = Inches(0.10)

    for i, dname in enumerate(header_names):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(dname)
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "000000")
        tcPr.append(shd)

    for wi, week in enumerate(grid):
        row = table.rows[wi + 2]
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(1.0)

        for di, cell_data in enumerate(week):
            cell = row.cells[di]
            cell.text = ""

            p_date = cell.paragraphs[0]
            p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_date.paragraph_format.space_before = Pt(0)
            p_date.paragraph_format.space_after = Pt(0)
            rd = p_date.add_run(f"{cal.month_abbr[cell_data['date'].month]} {cell_data['date'].day}")
            rd.bold = True
            rd.font.size = Pt(14)

            if cell_data["cycle_day"] is not None:
                p_day = cell.add_paragraph()
                p_day.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_day.paragraph_format.space_before = Pt(0)
                p_day.paragraph_format.space_after = Pt(0)
                rday = p_day.add_run(f"Day {cell_data['cycle_day']}")
                rday.italic = True
                rday.font.size = Pt(14)

                for lab in cell_data["labels"]:
                    p_lab = cell.add_paragraph()
                    p_lab.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_lab.paragraph_format.space_before = Pt(0)
                    p_lab.paragraph_format.space_after = Pt(0)
                    rl = p_lab.add_run(lab)
                    rl.font.size = Pt(14)
                    if lab.lower() != "rest":
                        rl.bold = True

    tbl_pr = table._element.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "auto")
        borders.append(e)
    tbl_pr.append(borders)

    doc.add_paragraph()

    for t in reg.therapies:
        route_phrase = _spell_route(t.route)
        verb = "Take" if t.route.strip().upper() == "PO" else "Given"
        freq_text = t.frequency.strip()
        dur_text = t.duration.strip()

        if t.total_doses is not None:
            total_doses = t.total_doses
        else:
            day_list = parse_day_spec(t.duration)
            total_doses = len(day_list) * _doses_per_day(t.frequency)

        dose_text = t.dose.strip()
        sentence = f"{t.name}: {verb} {dose_text} {route_phrase} {freq_text} on {dur_text} (total {total_doses} dose{'s' if total_doses != 1 else ''})."
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(sentence)
        run.font.size = Pt(12)

    doc.save(out_path)
    return True
