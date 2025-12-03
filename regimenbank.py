#!/usr/bin/env python3
from __future__ import annotations

import argparse
import calendar as cal
import datetime as dt
import os
import re
import sqlite3
import sys
import time
from dataclasses import dataclass, asdict, field, replace
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

# ---------------- config ----------------
SCHEMA_VERSION = 3
DEFAULT_DB = Path("regimenbank.db")
ROUTES = ["IV", "PO", "SQ", "IM", "IT"]

# ---------------- console style (notes only on selection) ----------------
def _supports_ansi() -> bool:
    return sys.stdout.isatty() and (
        os.name != "nt" or "WT_SESSION" in os.environ or "TERM" in os.environ
    )


def _italic(s: str) -> str:
    return f"\x1b[3m{s}\x1b[0m" if _supports_ansi() else s


# ---------------- models ----------------
@dataclass
class Chemotherapy:
    name: str
    route: str
    dose: str
    frequency: str  # FREE TEXT
    duration: str   # DAY MAP
    total_doses: Optional[int] = None  # NEW

    @staticmethod
    def from_dict(d: Dict[str, Any]) -> "Chemotherapy":
        return Chemotherapy(
            d["name"],
            d["route"],
            d["dose"],
            d["frequency"],
            d["duration"],
            d.get("total_doses"),  # NEW
        )


@dataclass
class Regimen:
    name: str
    disease_state: Optional[str] = None
    notes: Optional[str] = None  # stored; shown only on selection screen
    therapies: List[Chemotherapy] = field(default_factory=list)

    @staticmethod
    def from_dict(name: str, d: Dict[str, Any]) -> "Regimen":
        return Regimen(
            name=name,
            disease_state=d.get("disease_state"),
            notes=d.get("notes"),
            therapies=[Chemotherapy.from_dict(x) for x in d.get("therapies", [])],
        )

    def to_dict(self) -> Dict[str, Any]:
        return {
            "disease_state": self.disease_state,
            "notes": self.notes,
            "therapies": [asdict(t) for t in self.therapies],
        }

    def upsert_chemo(self, c: Chemotherapy) -> None:
        key = c.name.strip().lower()
        for i, t in enumerate(self.therapies):
            if t.name.strip().lower() == key:
                self.therapies[i] = c
                return
        self.therapies.append(c)


# ---------------- storage (SQLite) ----------------
class RegimenBank:
    """
    SQLite-backed regimen store.

    Tables:
      regimens(id, name, disease_state, notes, updated_at)
      therapies(
          id, regimen_id, name, route, dose,
          frequency, duration, total_doses
      )
    """

    def __init__(self, db_path: Path):
        self.db_path = db_path
        self.conn = sqlite3.connect(db_path, check_same_thread=False)
        self.conn.row_factory = sqlite3.Row
        # enforce FK constraints
        self.conn.execute("PRAGMA foreign_keys = ON")
        self._init_schema()

    def _init_schema(self) -> None:
        cur = self.conn.cursor()
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS regimens (
                id            INTEGER PRIMARY KEY AUTOINCREMENT,
                name          TEXT NOT NULL UNIQUE,
                disease_state TEXT,
                notes         TEXT,
                updated_at    TEXT NOT NULL
            )
        """
        )
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS therapies (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                regimen_id  INTEGER NOT NULL,
                name        TEXT NOT NULL,
                route       TEXT NOT NULL,
                dose        TEXT NOT NULL,
                frequency   TEXT NOT NULL,
                duration    TEXT NOT NULL,
                total_doses INTEGER,
                FOREIGN KEY (regimen_id) REFERENCES regimens(id) ON DELETE CASCADE
            )
        """
        )

        # --- lightweight migration: ensure total_doses exists ---
        cur.execute("PRAGMA table_info(therapies)")
        cols = [row["name"] for row in cur.fetchall()]
        if "total_doses" not in cols:
            cur.execute("ALTER TABLE therapies ADD COLUMN total_doses INTEGER")

        self.conn.commit()

    # ----- public API -----
    def list_regimens(self) -> List[str]:
        cur = self.conn.execute(
            "SELECT name FROM regimens ORDER BY name COLLATE NOCASE"
        )
        return [row["name"] for row in cur.fetchall()]

    def get_regimen(self, name: str) -> Optional[Regimen]:
        name = name.strip()
        cur = self.conn.execute(
            "SELECT id, name, disease_state, notes FROM regimens WHERE name = ?",
            (name,),
        )
        row = cur.fetchone()
        if not row:
            return None

        reg_id = row["id"]
        cur_t = self.conn.execute(
            "SELECT name, route, dose, frequency, duration, total_doses "
            "FROM therapies WHERE regimen_id = ? ORDER BY id",
            (reg_id,),
        )
        therapies = [
            Chemotherapy(
                trow["name"],
                trow["route"],
                trow["dose"],
                trow["frequency"],
                trow["duration"],
                trow["total_doses"],
            )
            for trow in cur_t.fetchall()
        ]

        return Regimen(
            name=row["name"],
            disease_state=row["disease_state"],
            notes=row["notes"],
            therapies=therapies,
        )

    def upsert_regimen(self, reg: Regimen) -> None:
        """
        Insert or update a regimen and fully replace its therapies.
        """
        now = time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
        with self.conn:
            # check if regimen exists
            cur = self.conn.execute(
                "SELECT id FROM regimens WHERE name = ?",
                (reg.name,),
            )
            row = cur.fetchone()
            if row:
                reg_id = row["id"]
                self.conn.execute(
                    "UPDATE regimens "
                    "SET disease_state = ?, notes = ?, updated_at = ? "
                    "WHERE id = ?",
                    (reg.disease_state, reg.notes, now, reg_id),
                )
                # replace therapies
                self.conn.execute(
                    "DELETE FROM therapies WHERE regimen_id = ?",
                    (reg_id,),
                )
            else:
                self.conn.execute(
                    "INSERT INTO regimens(name, disease_state, notes, updated_at) "
                    "VALUES(?, ?, ?, ?)",
                    (reg.name, reg.disease_state, reg.notes, now),
                )
                reg_id = self.conn.execute(
                    "SELECT id FROM regimens WHERE name = ?",
                    (reg.name,),
                ).fetchone()["id"]

            # insert therapies
            for t in reg.therapies:
                # auto-calc total_doses if missing
                total_doses = (
                    t.total_doses
                    if t.total_doses is not None
                    else len(parse_day_spec(t.duration))
                )
                self.conn.execute(
                    "INSERT INTO therapies("
                    "regimen_id, name, route, dose, frequency, duration, total_doses"
                    ") VALUES(?, ?, ?, ?, ?, ?, ?)",
                    (
                        reg_id,
                        t.name,
                        t.route,
                        t.dose,
                        t.frequency,
                        t.duration,
                        total_doses,
                    ),
                )

    def delete_regimen(self, name: str) -> bool:
        with self.conn:
            cur = self.conn.execute(
                "DELETE FROM regimens WHERE name = ?",
                (name.strip(),),
            )
            return cur.rowcount > 0

    def save_as(self, reg: Regimen, new_name: str) -> None:
        r2 = replace(reg, name=new_name)
        self.upsert_regimen(r2)

    def close(self) -> None:
        try:
            self.conn.close()
        except Exception:
            pass


# ---------------- small IO helpers ----------------
def _choose(prompt: str, options: List[str], allow_new=False) -> Tuple[str, bool]:
    print(f"\n{prompt}")
    for i, o in enumerate(options, 1):
        print(f"  {i}. {o}")
    if allow_new:
        print("  n. <Add new>")
    while True:
        s = input("Select: ").strip().lower()
        if allow_new and s == "n":
            val = input("Enter new name: ").strip()
            if val:
                return val, True
        if s.isdigit():
            k = int(s)
            if 1 <= k <= len(options):
                return options[k - 1], False
        print("Invalid. Try again.")


def _req(label: str, pre: Optional[str] = None) -> str:
    while True:
        s = input(f"{label}{f' [{pre}]' if pre else ''}: ").strip()
        if s:
            return s
        if pre:
            return pre
        print("Required.")


def _opt(label: str, pre: Optional[str] = None) -> Optional[str]:
    s = input(f"{label}{f' [{pre}]' if pre else ''} (optional): ").strip()
    return s or pre


def _parse_date(label: str, default: Optional[dt.date] = None) -> dt.date:
    while True:
        hint = f" [{default.strftime('%m/%d/%y')}]" if default else ""
        s = input(f"{label}{hint}: ").strip().lower()
        if not s and default:
            return default
        if s in {"t", "today"}:
            return dt.date.today()
        if s.startswith("+") and s[1:].isdigit():
            return dt.date.today() + dt.timedelta(days=int(s[1:]))
        for fmt in ("%Y-%m-%d", "%m/%d/%y", "%m/%d/%Y"):
            try:
                return dt.datetime.strptime(s, fmt).date()
            except ValueError:
                pass
        print("Use YYYY-MM-DD, M/D/YY, M/D/YYYY, 'today', or +N.")


def _edit_agent(base: Chemotherapy) -> Chemotherapy:
    name = _req("Agent name", base.name)
    route = _req(f"Route ({'/'.join(ROUTES)})", base.route)
    dose  = _req("Dose", base.dose)
    frequency = _req(
        "Frequency (free text: once / daily / BID / TID / weekly ...)",
        base.frequency,
    )
    duration  = _req(
        "Day map for calendar (e.g., 'Days 1–7' or 'Days 1,8,15')",
        base.duration,
    )

    # ------ NEW: optional total_doses override ------
    # Show existing total_doses if present
    td_default = "" if base.total_doses is None else str(base.total_doses)
    td_raw = input(
        f"Total doses (optional; leave blank to auto-calc from days)"
        f"{f' [{td_default}]' if td_default else ''}: "
    ).strip()

    if not td_raw and td_default:
        # keep previous value
        try:
            td_val: Optional[int] = int(td_default)
        except ValueError:
            td_val = None
    elif not td_raw:
        td_val = None
    else:
        try:
            td_val = int(td_raw)
        except ValueError:
            print("Could not parse total doses as an integer; will auto-calc from duration.")
            td_val = None

    return Chemotherapy(name, route, dose, frequency, duration, td_val)


# ---------------- notes on selection only ----------------
def _select_regimen_with_notes(
    bank: RegimenBank, title: str, allow_new=False
) -> Tuple[Optional[str], bool]:
    names = bank.list_regimens()
    if not names and not allow_new:
        return None, False
    if not names and allow_new:
        val = input("No regimens yet. Enter a new regimen name: ").strip()
        return (val if val else None), True

    print(f"\n{title}")
    index_map: Dict[str, str] = {}
    for i, n in enumerate(names, 1):
        reg = bank.get_regimen(n)
        note = f"  {_italic('— ' + reg.notes)}" if reg and reg.notes else ""
        print(f"  {i}. {n}{note}")
        index_map[str(i)] = n
    if allow_new:
        print("  n. <Add new>")

    while True:
        sel = input("Select: ").strip().lower()
        if allow_new and sel == "n":
            val = input("Enter new regimen name: ").strip()
            if val:
                return val, True
        if sel in index_map:
            return index_map[sel], False
        print("Invalid. Try again.")


# ---------------- day-map parsing ----------------
def parse_day_spec(day_spec: str) -> List[int]:
    """
    Parse a day specification into explicit day numbers.

    Accepts flexible formats like:
      - 'Days 1-7'
      - 'Day 1-7'
      - 'Days: 1-3, 5, 7-9'
      - '1-7'
      - '1, 4, 8'
    """
    if not day_spec:
        return []

    s = day_spec.replace("–", "-").strip().lower()

    # Strip leading 'day' / 'days' and optional punctuation
    s = re.sub(r"^days?\s*[:\-]?\s*", "", s)

    if not s:
        return []

    tokens = re.split(r"[,\s]+", s)
    out: List[int] = []

    for tok in tokens:
        if not tok:
            continue
        if "-" in tok:
            try:
                a_str, b_str = tok.split("-", 1)
                a, b = int(a_str), int(b_str)
                if a <= b:
                    out.extend(range(a, b + 1))
            except ValueError:
                # Skip malformed segments instead of killing everything
                continue
        else:
            try:
                out.append(int(tok))
            except ValueError:
                continue

    return sorted(set(d for d in out if d >= 1))


# ---------------- grid + exports ----------------
def compute_calendar_grid(reg: Regimen, start: dt.date, cycle_len: int):
    max_day = cycle_len
    by_day: Dict[int, List[str]] = {d: [] for d in range(1, cycle_len + 1)}
    for t in reg.therapies:
        dlist = [d for d in parse_day_spec(t.duration) if d <= cycle_len]
        if dlist:
            max_day = max(max_day, max(dlist))
            for d in dlist:
                by_day.setdefault(d, []).append(t.name)

    # first Sunday before/at start
    first_sun = start - dt.timedelta(days=(start.weekday() + 1) % 7)
    last_needed = start + dt.timedelta(days=max_day - 1)
    # last Saturday on/after last_needed  (weekday: Mon=0..Sun=6, Sat=5)
    last_sat = last_needed + dt.timedelta(days=(5 - last_needed.weekday()) % 7)

    grid: List[List[Dict[str, Any]]] = []
    d = first_sun
    week: List[Dict[str, Any]] = []
    while d <= last_sat:
        entry: Dict[str, Any] = {"date": d, "cycle_day": None, "labels": []}
        if d >= start:
            cd = (d - start).days + 1
            if 1 <= cd <= max_day:
                entry["cycle_day"] = cd
                entry["labels"] = by_day.get(cd, []) or ["Rest"]
        week.append(entry)
        if len(week) == 7:
            grid.append(week)
            week = []
        d += dt.timedelta(days=1)
    if week:
        while len(week) < 7:
            week.append({"date": d, "cycle_day": None, "labels": []})
            d += dt.timedelta(days=1)
        grid.append(week)
    return first_sun, last_sat, max_day, grid


def make_calendar_text(
    reg: Regimen,
    start: dt.date,
    cycle_len: int,
    cycle_label: str,
    note: Optional[str] = None,
) -> str:
    first_sun, last_sat, _, grid = compute_calendar_grid(reg, start, cycle_len)
    months = cal.month_name[first_sun.month]
    if first_sun.month != last_sat.month or first_sun.year != last_sat.year:
        months += f" - {cal.month_name[last_sat.month]}"
    year = (
        str(first_sun.year)
        if first_sun.year == last_sat.year
        else f"{first_sun.year}-{last_sat.year}"
    )

    lines = [f"{reg.name} — {cycle_label}", f"{months} {year}"]
    if note:
        lines.append(f"Note: {note}")
    lines.append("Sun       Mon       Tue       Wed       Thu       Fri       Sat")

    colw = 12
    for week in grid:
        cells: List[List[str]] = []
        maxh = 0
        for c in week:
            block = [f"{cal.month_abbr[c['date'].month]} {c['date'].day}"]
            if c["cycle_day"] is not None:
                block.append(f"Day {c['cycle_day']}")
                block.extend(c["labels"])
            maxh = max(maxh, len(block))
            cells.append(block)
        for r in range(maxh):
            row = []
            for b in cells:
                txt = b[r] if r < len(b) else ""
                row.append(txt.ljust(colw))
            lines.append(" ".join(row))
        lines.append("")
    return "\n".join(lines)


def _spell_route(route: str) -> str:
    r = route.strip().upper()
    mapping = {
        "PO": "by mouth",
        "IV": "intravenously",
        "SQ": "subcutanously",
        "IT": "intrathecally. Given during lumbar puncture",
    }
    return mapping.get(r, route)  # fall back to raw text if unknown


def export_calendar_docx(
    reg: Regimen,
    start: dt.date,
    cycle_len: int,
    out_path: Path,
    cycle_label: str,
    note: Optional[str] = None,
) -> bool:
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return False

    first_sun, last_sat, _, grid = compute_calendar_grid(reg, start, cycle_len)
    months = cal.month_name[first_sun.month]
    if first_sun.month != last_sat.month or first_sun.year != last_sat.year:
        months += f" - {cal.month_name[last_sat.month]}"
    year = (
        str(first_sun.year)
        if first_sun.year == last_sat.year
        else f"{first_sun.year}-{last_sat.year}"
    )

    doc = Document()
    # Set default document font to Calibri
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    style.font.size = Pt(12)   # optional default size

    sec = doc.sections[0]

    # Page layout: landscape, 0.5" margins
    sec.orientation = 1
    sec.page_width, sec.page_height = Inches(11), Inches(8.5)
    sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Inches(
        0.5
    )

    # ---------- HEADER: Name/DOB left, logo right ----------
    hdr = sec.header
    try:
        htbl = hdr.add_table(rows=1, cols=2, width=sec.page_width)
    except TypeError:
        htbl = hdr.add_table(rows=1, cols=2)
    htbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Left cell: name + DOB, italic, small
    left_cell = htbl.cell(0, 0)
    left_cell.text = ""

    p_name = left_cell.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(0)
    r_name = p_name.add_run("First Last")
    r_name.italic = True
    r_name.font.size = Pt(10)

    p_dob = left_cell.add_paragraph()
    p_dob.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_dob.paragraph_format.space_before = Pt(0)
    p_dob.paragraph_format.space_after = Pt(0)
    r_dob = p_dob.add_run("DOB: M/DD/YYYY")
    r_dob.italic = True
    r_dob.font.size = Pt(10)

    # remove any default empty paragraph if present
    if left_cell.paragraphs and left_cell.paragraphs[0].text == "":
        try:
            left_cell._element.remove(left_cell.paragraphs[0]._element)
        except Exception:
            pass

    # Right cell: UCM logo aligned right (if present)
    right_cell = htbl.cell(0, 1)
    right_p = right_cell.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    logo = Path("ucm.png")
    if logo.exists():
        try:
            right_p.add_run().add_picture(str(logo), height=Inches(0.76))
        except Exception as e:
            print(f"[WARN] Could not insert logo: {e}")

    # ---------- CALENDAR TABLE ----------
    rows = len(grid) + 2  # row 0 = title; row 1 = weekday header
    cols = 7
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # ----- Row 0: merged title row -----
    title_row = table.rows[0]
    title_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    title_row.height = Inches(0.72)

    first_cell = title_row.cells[0]
    for j in range(1, cols):
        first_cell.merge(title_row.cells[j])
    c = first_cell
    c.text = ""

    # Line 1: "Chemotherapy Calendar"
    p1 = c.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run("Chemotherapy Calendar")
    r1.bold = True
    r1.font.size = Pt(14)

    # Line 2: regimen name - Cycle X   (cycle bold)
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)

    r2a = p2.add_run(reg.name + " - ")
    r2a.font.size = Pt(14)

    r2b = p2.add_run(cycle_label)
    r2b.bold = True
    r2b.font.size = Pt(14)

    # Line 3: month/year range
    p3 = c.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(0)
    r3 = p3.add_run(f"{months} {year}")
    r3.font.size = Pt(14)

    # Optional note line
    if note:
        p4 = c.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.paragraph_format.space_before = Pt(0)
        p4.paragraph_format.space_after = Pt(0)
        r4 = p4.add_run(note)
        r4.italic = True
        r4.font.size = Pt(11)

    # ----- Row 1: weekday header row (black with white text) -----
    header_names = [
        "Sunday",
        "Monday",
        "Tuesday",
        "Wednesday",
        "Thursday",
        "Friday",
        "Saturday",
    ]
    header_row = table.rows[1]
    header_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    header_row.height = Inches(0.10)

    for i, dname in enumerate(header_names):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(dname)
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # black shading
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "000000")
        tcPr.append(shd)

    # ----- Rows 2+: calendar weeks -----
    for wi, week in enumerate(grid):
        row = table.rows[wi + 2]
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(1.0)

        for di, cell_data in enumerate(week):
            cell = row.cells[di]
            cell.text = ""

            # Date – right aligned
            p_date = cell.paragraphs[0]
            p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_date.paragraph_format.space_before = Pt(0)
            p_date.paragraph_format.space_after = Pt(0)
            rd = p_date.add_run(
                f"{cal.month_abbr[cell_data['date'].month]} {cell_data['date'].day}"
            )
            rd.bold = True
            rd.font.size = Pt(14)

            if cell_data["cycle_day"] is not None:
                # "Day X" – left aligned
                p_day = cell.add_paragraph()
                p_day.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_day.paragraph_format.space_before = Pt(0)
                p_day.paragraph_format.space_after = Pt(0)
                rday = p_day.add_run(f"Day {cell_data['cycle_day']}")
                rday.italic = True
                rday.font.size = Pt(14)

                # Med labels – left aligned, bold when not "rest"
                for lab in cell_data["labels"]:
                    p_lab = cell.add_paragraph()
                    p_lab.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_lab.paragraph_format.space_before = Pt(0)
                    p_lab.paragraph_format.space_after = Pt(0)
                    rl = p_lab.add_run(lab)
                    rl.font.size = Pt(14)
                    if lab.lower() != "rest":
                        rl.bold = True

    # ----- Table borders (thin) -----
    tbl_pr = table._element.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "auto")
        borders.append(e)
    tbl_pr.append(borders)

    # ---------- Patient-friendly bullets below calendar ----------
    doc.add_paragraph()

    for t in reg.therapies:
        route_phrase = _spell_route(t.route)
        verb = "Take" if t.route.strip().upper() == "PO" else "Given"

        freq_text = t.frequency.strip()
        dur_text = t.duration.strip()

        if t.total_doses is not None:
            total_doses = t.total_doses
        else:
            day_list = parse_day_spec(t.duration)
            total_doses = len(day_list)

        sentence = (
            f"{t.name}: "
            f"{verb} {route_phrase} "
            f"{freq_text} "
            f"on {dur_text} "
            f"(total {total_doses} dose{'s' if total_doses != 1 else ''})."
        )

        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(sentence)
        run.font.size = Pt(11)

    doc.save(out_path)
    return True



# ---------------- catalog + editors ----------------
def build_agent_catalog(bank: RegimenBank) -> Dict[str, List[Chemotherapy]]:
    cat: Dict[str, List[Chemotherapy]] = {}
    for rn in bank.list_regimens():
        r = bank.get_regimen(rn)
        if not r:
            continue
        for t in r.therapies:
            key = t.name.strip().lower()
            cat.setdefault(key, [])
            if not any(
                v.route == t.route
                and v.dose == t.dose
                and v.frequency == t.frequency
                and v.duration == t.duration
                for v in cat[key]
            ):
                cat[key].append(t)
    return cat


def choose_agent_from_catalog(cat: Dict[str, List[Chemotherapy]]) -> Optional[Chemotherapy]:
    if not cat:
        print("No saved agents yet.")
        return None
    names = sorted(cat.keys())
    disp = [n.title() for n in names]
    print("\nChoose agent name:")
    for i, d in enumerate(disp, 1):
        print(f"  {i}. {d}")
    while True:
        s = input("Select: ").strip()
        if s.isdigit() and 1 <= int(s) <= len(disp):
            key = names[int(s) - 1]
            vs = cat[key]
            break
        print("Invalid.")
    if len(vs) == 1:
        return replace(vs[0])
    for i, v in enumerate(vs, 1):
        print(f"  {i}. {v.route} | {v.dose} | {v.frequency} | {v.duration}")
    while True:
        s = input("Pick variant #: ").strip()
        if s.isdigit() and 1 <= int(s) <= len(vs):
            return replace(vs[int(s) - 1])
        print("Invalid.")


# ---------------- unified save-or-save-as ----------------
def _persist_menu(bank: RegimenBank, reg: Regimen) -> None:
    """
    Offer: 1) Save As (new name), 2) Save (same name), 3) Don't save.
    Overwrites only happen if the user explicitly types 'yes'.
    """
    exists = bank.get_regimen(reg.name) is not None

    print("\nPersist changes:")
    if exists:
        print(f"  Current name '{reg.name}' already exists in regimen bank.")
    else:
        print(f"  Current name '{reg.name}' does not yet exist in regimen bank.")

    # Keep the safe option first
    print("  1. Save As (create NEW regimen with a new name)")
    if exists:
        print(f"  2. Save (OVERWRITE existing '{reg.name}')")
    else:
        print("  2. Save (create using current name)")
    print("  3. Don't save")

    while True:
        sel = input("Select [1-3]: ").strip()
        if sel == "1":
            # Always ask for a name for Save As
            while True:
                new_name = input("New regimen name: ").strip()
                if not new_name:
                    print("Name required.")
                    continue

                new_exists = bank.get_regimen(new_name) is not None
                if new_exists:
                    print(f"A regimen named '{new_name}' already exists.")
                    confirm = input(
                        "Overwrite it? Type 'yes' to confirm, or anything else to cancel: "
                    ).strip().lower()
                    if confirm != "yes":
                        print("Overwrite cancelled. Choose a different name.")
                        continue

                bank.save_as(reg, new_name)
                print(f"Saved as '{new_name}'.")
                return

        elif sel == "2":
            if exists:
                print(f"WARNING: This will overwrite existing regimen '{reg.name}'.")
                confirm = input(
                    "Type 'yes' to overwrite, or anything else to cancel: "
                ).strip().lower()
                if confirm != "yes":
                    print("Overwrite cancelled.")
                    continue
            bank.upsert_regimen(reg)
            print(f"Saved '{reg.name}'.")
            return

        elif sel == "3":
            print("Not saved to regimen bank.")
            return

        else:
            print("Choose 1–3.")


# ---------------- wizard (no scaffolds) ----------------
def wizard(bank: RegimenBank) -> None:
    rname, is_new = _select_regimen_with_notes(
        bank, "Select a regimen or add a new one:", allow_new=True
    )
    if rname is None:
        print("No regimen selected.")
        return
    reg = bank.get_regimen(rname) if not is_new else Regimen(name=rname)

    if is_new or reg is None:
        print(f"\nCreating NEW regimen '{rname}'.")
        reg = reg or Regimen(name=rname)
    else:
        print(f"\nEditing EXISTING regimen '{rname}'.")
        print(
            "Hint: Use 'Persist → Save As' if you want to create a new regimen (e.g., 7+3 → 7+3+ven)."
        )

    reg.disease_state = _opt("Disease state", reg.disease_state)
    reg.notes = _opt("Regimen notes (selection aid only)", reg.notes)

    while True:
        print("\nTherapies:")
        if not reg.therapies:
            print("  (none)")
        else:
            for i, t in enumerate(reg.therapies, 1):
                print(
                    f"  {i}. {t.name} | {t.route} | {t.dose} | {t.frequency} | {t.duration}"
                )
        print(
            "\nActions: 1) Add new 2) Add from existing library  3) Edit  "
            "4) Remove agent  5) Persist (Save / Save As)  6) Finish"
        )
        ch = input("Select [1-6]: ").strip()
        if ch == "1":
            reg.upsert_chemo(Chemotherapy("", "", "", "", ""))
            reg.therapies[-1] = _edit_agent(reg.therapies[-1])
        elif ch == "2":
            tmpl = choose_agent_from_catalog(build_agent_catalog(bank))
            if tmpl:
                reg.upsert_chemo(_edit_agent(tmpl))
        elif ch == "3":
            if not reg.therapies:
                print("No agents.")
                continue
            k = input("Agent # to edit: ").strip()
            if k.isdigit() and 1 <= int(k) <= len(reg.therapies):
                reg.therapies[int(k) - 1] = _edit_agent(reg.therapies[int(k) - 1])
            else:
                print("Invalid.")
        elif ch == "4":
            if not reg.therapies:
                print("No agents.")
                continue
            k = input("Agent # to remove: ").strip()
            if k.isdigit() and 1 <= int(k) <= len(reg.therapies):
                gone = reg.therapies.pop(int(k) - 1)
                print(f"Removed {gone.name}.")
            else:
                print("Invalid.")
        elif ch == "5":
            _persist_menu(bank, reg)
        elif ch == "6":
            if input("Save before finishing? [y/N]: ").strip().lower() == "y":
                _persist_menu(bank, reg)
            print("Done.")
            return
        else:
            print("Choose 1–6.")


# ---------------- prep editor (no notes UI; save/save-as) ----------------
def prep_editor(base: Regimen, bank: RegimenBank) -> Tuple[Regimen, Optional[str]]:
    # Work on a copy so we never mutate the base regimen directly
    work = Regimen(
        base.name,
        base.disease_state,
        base.notes,
        [replace(t) for t in base.therapies],
    )
    print(f"\n=== Calendar Prep Editor for '{work.name}' ===")
    # No calendar-specific note prompt anymore
    inst_note: Optional[str] = None

    while True:
        print("\nTherapies:")
        if not work.therapies:
            print("  (none)")
        else:
            for i, t in enumerate(work.therapies, 1):
                print(
                    f"  {i}. {t.name} | {t.route} | {t.dose} | {t.frequency} | {t.duration}"
                )

        choice = input(
            "\nPress Enter to continue with this calendar copy, "
            "or type 'e' to edit: "
        ).strip().lower()

        if choice == "":
            # Just use the regimen as-is for this calendar; no edits, no save prompts
            return work, inst_note

        if choice not in {"e", "1"}:
            print("Type 'e' to edit, or just press Enter to continue.")
            continue

        # --- Edit sub-menu ---
        changed = False
        while True:
            print("\nEdit actions:")
            print("  1) Edit existing agent")
            print("  2) Add new agent")
            print("  3) Add from existing catalog")
            print("  4) Remove agent")
            print("  5) Done editing")

            ch = input("Select [1-5] (or press Enter for 'Done editing'): ").strip()

            if ch == "" or ch == "5":
                break

            elif ch == "1":
                if not work.therapies:
                    print("No agents to edit.")
                    continue
                k = input("Agent # to edit: ").strip()
                if k.isdigit() and 1 <= int(k) <= len(work.therapies):
                    work.therapies[int(k) - 1] = _edit_agent(
                        work.therapies[int(k) - 1]
                    )
                    changed = True
                else:
                    print("Invalid.")

            elif ch == "2":
                work.upsert_chemo(Chemotherapy("", "", "", "", ""))
                work.therapies[-1] = _edit_agent(work.therapies[-1])
                changed = True

            elif ch == "3":
                tmpl = choose_agent_from_catalog(build_agent_catalog(bank))
                if tmpl:
                    work.upsert_chemo(_edit_agent(tmpl))
                    changed = True

            elif ch == "4":
                if not work.therapies:
                    print("No agents to remove.")
                    continue
                k = input("Agent # to remove: ").strip()
                if k.isdigit() and 1 <= int(k) <= len(work.therapies):
                    gone = work.therapies.pop(int(k) - 1)
                    print(f"Removed {gone.name}.")
                    changed = True
                else:
                    print("Invalid.")

            else:
                print("Choose 1–5.")

        # After editing session, optionally persist back to regimen bank
        if changed:
            ans = input(
                "Save these edits back to the regimen bank (Save / Save As)? [y/N]: "
            ).strip().lower()
            if ans == "y":
                _persist_menu(bank, work)
        return work, inst_note


# ---------------- labels ----------------
def _cycle_label() -> str:
    print("\nPhase: 1) Cycle #  2) Induction")
    while True:
        s = input("Select [1-2]: ").strip()
        if s == "1":
            while True:
                n = input("Cycle number: ").strip()
                if n.isdigit() and int(n) >= 1:
                    return f"Cycle {int(n)}"
                print("Positive integer.")
        if s == "2":
            return "Induction"
        print("Choose 1 or 2.")


# ---------------- calendar flow ----------------
def calendar_flow(bank: RegimenBank) -> None:
    names = bank.list_regimens()
    if not names:
        print("No regimens yet → launching wizard.")
        wizard(bank)
        names = bank.list_regimens()
        if not names:
            print("Still empty. Exiting.")
            return

    rname, _ = _select_regimen_with_notes(bank, "Select a regimen:", allow_new=False)
    if not rname:
        print("No regimen selected.")
        return
    base = bank.get_regimen(rname)
    if not base:
        print("Not found.")
        return

    reg, note = prep_editor(base, bank)
    start = _parse_date("Cycle start date", default=dt.date.today())

    while True:
        s = input("Cycle length in days [28]: ").strip()
        if not s:
            cycle = 28
            break
        if s.isdigit() and int(s) >= 1:
            cycle = int(s)
            break
        print("Positive integer.")

    label = _cycle_label()

    txt = make_calendar_text(reg, start, cycle, label, note)
    print("\n" + txt + "\n")

    safe = "".join(c if c.isalnum() or c in ("-", "_") else "_" for c in reg.name)
    if not safe:
        safe = "regimen"
    basefn = f"{safe}_{label.replace(' ', '').lower()}_{start.isoformat()}"
    if input(f"Save text [{basefn}.txt]? [Y/n]: ").strip().lower() != "n":
        Path(f"{basefn}.txt").write_text(txt, encoding="utf-8")
        print(f"Saved {basefn}.txt")

    if input(f"Export DOCX [{basefn}.docx]? [y/N]: ").strip().lower() == "y":
        ok = export_calendar_docx(reg, start, cycle, Path(f"{basefn}.docx"), label, note)
        if ok:
            print(f"Saved {basefn}.docx")
        else:
            print("Install python-docx:  pip install python-docx")
            print("Then re-run calendar export.")


# ---------------- CLI ----------------
def _parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        description="Chemotherapy regimen bank + calendar (TXT/DOCX)"
    )
    p.add_argument(
        "--db",
        type=Path,
        default=DEFAULT_DB,
        help="Path to regimenbank.db (SQLite)",
    )
    sub = p.add_subparsers(dest="cmd", required=True)
    sub.add_parser("wizard", help="Create or edit regimens")
    sub.add_parser("calendar", help="Prep + generate calendar")
    sub.add_parser("list", help="List regimens")
    sh = sub.add_parser("show", help="Show regimen")
    sh.add_argument("--name", required=True)
    dr = sub.add_parser("delete-regimen", help="Delete regimen")
    dr.add_argument("--name", required=True)
    return p


def _print_regimen(r: Regimen) -> None:
    print(f"\nRegimen: {r.name}")
    if r.disease_state:
        print(f"Disease State: {r.disease_state}")
    if r.notes:
        print(f"Notes: {r.notes}")
    if not r.therapies:
        print("Therapies: (none)")
    else:
        print("Therapies:")
        for i, t in enumerate(r.therapies, 1):
            print(
                f"  {i}. {t.name} | {t.route} | {t.dose} | {t.frequency} | {t.duration} "
                f"(total_doses={t.total_doses})"
            )
    print("")


def main(argv: List[str]) -> int:
    args = _parser().parse_args(argv)
    bank = RegimenBank(args.db)
    try:
        if args.cmd == "wizard":
            wizard(bank)
            return 0
        if args.cmd == "calendar":
            calendar_flow(bank)
            return 0
        if args.cmd == "list":
            xs = bank.list_regimens()
            print("(no regimens)" if not xs else "\n".join(xs))
            return 0
        if args.cmd == "show":
            r = bank.get_regimen(args.name)
            if not r:
                print("Not found.")
                return 1
            _print_regimen(r)
            return 0
        if args.cmd == "delete-regimen":
            ok = bank.delete_regimen(args.name)
            print("Deleted." if ok else "Not found.")
            return 0 if ok else 1
        return 1
    finally:
        bank.close()


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
