# app.py — Streamlit UI
from pathlib import Path
from io import BytesIO
import datetime as dt
import streamlit as st

# Reuse your existing core logic
import regimenbank as rb

DB_PATH = Path("regimenbank.json")

st.set_page_config(page_title="Chemo Calendar", layout="wide")

def _rerun():
    # Streamlit >= 1.27 uses st.rerun(); older builds used experimental_rerun
    import streamlit as st
    if hasattr(st, "rerun"):
        st.rerun()
    else:
        st.experimental_rerun()

# ---------- Utils ----------
def load_bank() -> rb.RegimenBank:
    return rb.RegimenBank(DB_PATH)

def get_regimens(bank: rb.RegimenBank):
    return bank.list_regimens()

def export_docx_bytes(reg: rb.Regimen, start: dt.date, cycle_len: int, cycle_label: str) -> bytes:
    """
    Same export as in regimenbank, but to bytes (no file i/o).
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except Exception:
        st.warning("python-docx is not installed. Run: pip install python-docx")
        return b""

    first_sun, last_sat, max_day, grid = rb.compute_calendar_grid(reg, start, cycle_len)

    import calendar
    months = calendar.month_name[first_sun.month]
    if first_sun.month != last_sat.month or first_sun.year != last_sat.year:
        months += f" - {calendar.month_name[last_sat.month]}"
    title_year = str(first_sun.year) if first_sun.year == last_sat.year else f"{first_sun.year}-{last_sat.year}"

    doc = Document()
    section = doc.sections[0]
    section.orientation = 1
    section.page_width, section.page_height = Inches(11), Inches(8.5)
    section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Inches(0.5)

    table = doc.add_table(rows=len(grid) + 3, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    def merge_row(row_idx: int):
        row = table.rows[row_idx]
        first_cell = row.cells[0]
        for j in range(1, 7):
            first_cell.merge(row.cells[j])
        return first_cell

    # Row 0: Title
    c0 = merge_row(0)
    p = c0.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Chemotherapy Calendar\n"); r.bold = True; r.font.size = Pt(14)
    r2 = p.add_run(f"{reg.name}  - {cycle_label}\n"); r2.font.size = Pt(12)
    r3 = p.add_run(f"{months} {title_year}"); r3.font.size = Pt(12)

    # Row 1: Name/DOB blanks
    c1 = merge_row(1)
    p = c1.paragraphs[0]
    p.add_run("Patient Name: ").bold = True; p.add_run("__________________________    ")
    p.add_run("DOB: ").bold = True; p.add_run("______________")

    # Row 2: Weekday header
    hdr = ["Sunday","Monday","Tuesday","Wednesday","Thursday","Friday","Saturday"]
    for i, text in enumerate(hdr):
        cell = table.cell(2, i)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        x = para.add_run(text); x.bold = True; x.font.size = Pt(10)

    # Body rows
    body_start = 3
    import calendar as calmod
    for wi, week in enumerate(grid):
        for di, cell in enumerate(week):
            c = table.cell(body_start + wi, di)
            para = c.paragraphs[0]
            r1 = para.add_run(f"{calmod.month_abbr[cell['date'].month]} {cell['date'].day}\n")
            r1.bold = True; r1.font.size = Pt(9)
            if cell["cycle_day"] is not None:
                r2 = para.add_run(f"Day {cell['cycle_day']}\n"); r2.font.size = Pt(9)
                if cell["labels"]:
                    for lab in cell["labels"]:
                        para.add_run(f"{lab}\n").font.size = Pt(9)

    # Borders
    def set_borders(tbl):
        tbl_pr = tbl._element.tblPr
        tbl_borders = OxmlElement('w:tblBorders')
        for edge in ('top','left','bottom','right','insideH','insideV'):
            e = OxmlElement(f'w:{edge}')
            e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:space'),'0'); e.set(qn('w:color'),'auto')
            tbl_borders.append(e)
        tbl_pr.append(tbl_borders)
    set_borders(table)

    doc.add_paragraph().add_run("Note: Add patient identifiers (Name/DOB) after generating this document.").italic = True

    buff = BytesIO()
    doc.save(buff)
    buff.seek(0)
    return buff.read()

def cycle_label_ui() -> str:
    choice = st.radio("Phase", ["Cycle #", "Induction"], horizontal=True)
    if choice == "Induction":
        return "Induction"
    n = st.number_input("Cycle number", min_value=1, step=1, value=1)
    return f"Cycle {int(n)}"

def regimen_select_ui(bank: rb.RegimenBank, key: str = "regsel") -> str | None:
    names = get_regimens(bank)
    if not names:
        st.info("No regimens yet. Create one below.")
        return None
    idx = st.selectbox("Choose regimen", options=range(len(names)), format_func=lambda i: names[i], key=key)
    return names[idx]

def agent_catalog_options(bank: rb.RegimenBank):
    cat = rb.build_agent_catalog(bank)
    name_keys = sorted(cat.keys())
    display_names = [n.title() for n in name_keys]
    return cat, name_keys, display_names

# ---------- Sidebar ----------
with st.sidebar:
    st.title("Chemo Calendar")
    view = st.radio("Mode", ["Regimens", "Calendar"], horizontal=False)

bank = load_bank()

# ---------- Regimens View ----------
if view == "Regimens":
    st.header("Regimens")

    colA, colB = st.columns([2,1])

    with colA:
        # existing or new
        names = get_regimens(bank)
        create_new = st.toggle("Create new regimen")
        if create_new:
            reg_name = st.text_input("Regimen name (e.g., AZA/VEN 70 mg)")
            disease_state = st.text_input("Disease state (optional)")
            do_scaffold = st.checkbox("Quick scaffold AZA/VEN", value=False)
            if st.button("Create regimen", type="primary", disabled=not reg_name.strip()):
                reg = rb.Regimen(name=reg_name.strip(), disease_state=(disease_state.strip() or None))
                if do_scaffold and "ven" in reg_name.lower():
                    aza_dose = st.session_state.get("azadosetmp") or "75 mg/m^2"
                    ven_dose = st.session_state.get("vendosetmp") or "70 mg"
                    ven_days = st.session_state.get("vendaystmp") or 21
                    reg.upsert_chemo(rb.Chemotherapy("Azacitidine","IV", aza_dose, "Days 1–7","7 days"))
                    reg.upsert_chemo(rb.Chemotherapy("Venetoclax","PO", ven_dose, f"Days 1–{ven_days}", f"{ven_days} days"))
                bank.upsert_regimen(reg)
                st.success(f"Created '{reg.name}'")
                st._rerun()
        else:
            if not names:
                st.info("No regimens yet. Toggle 'Create new regimen'.")
            else:
                sel = regimen_select_ui(bank, key="reg_edit_sel")
                if sel:
                    reg = bank.get_regimen(sel)
                    st.subheader(sel)
                    disease_state = st.text_input("Disease state (optional)", value=reg.disease_state or "")
                    if st.button("Save disease state"):
                        reg.disease_state = disease_state.strip() or None
                        bank.upsert_regimen(reg)
                        st.success("Saved.")
                    st.markdown("### Agents")
                    # table-style listing
                    if not reg.therapies:
                        st.caption("No agents yet.")
                    else:
                        for i, t in enumerate(reg.therapies):
                            with st.container(border=True):
                                st.write(f"**{t.name}** — {t.route} • {t.dose}")
                                st.write(f"Freq: {t.frequency} • Duration: {t.duration}")
                                c1, c2 = st.columns(2)
                                if c1.button(f"Edit #{i+1}", key=f"edit{i}"):
                                    st.session_state[f"edit_idx"] = i
                                if c2.button(f"Remove #{i+1}", key=f"rem{i}"):
                                    reg.therapies.pop(i)
                                    bank.upsert_regimen(reg)
                                    st.warning("Removed.")
                                    _rerun()
                        # inline editor
                        edit_idx = st.session_state.get("edit_idx", None)
                        if isinstance(edit_idx, int) and 0 <= edit_idx < len(reg.therapies):
                            st.markdown("#### Edit agent")
                            t = reg.therapies[edit_idx]
                            name = st.text_input("Agent name", value=t.name, key="e_name")
                            route = st.selectbox("Route", ["IV","PO","SQ","IM","IT",], index=["IV","PO","SQ","IM","IT"].index(t.route) if t.route in ["IV","PO","SQ","IM","IT"] else 0)
                            dose = st.text_input("Dose", value=t.dose, key="e_dose")
                            freq = st.text_input("Frequency", value=t.frequency, key="e_freq")
                            dur  = st.text_input("Duration", value=t.duration, key="e_dur")
                            if st.button("Save agent"):
                                reg.therapies[edit_idx] = rb.Chemotherapy(name, route, dose, freq, dur)
                                bank.upsert_regimen(reg)
                                st.success("Updated.")
                                st.session_state.pop("edit_idx", None)
                                _rerun()

                    st.markdown("### Add agent")
                    tab1, tab2 = st.tabs(["New", "From existing"])
                    with tab1:
                        n = st.text_input("Agent name", key="n_name")
                        r = st.selectbox("Route", ["IV","PO","SQ","IM","IT"], key="n_route")
                        d = st.text_input("Dose", key="n_dose")
                        f = st.text_input("Frequency (e.g., Days 1–7 or Days 1,8,15)", key="n_freq")
                        u = st.text_input("Duration (e.g., 7 days)", key="n_dur")
                        if st.button("Add new agent", disabled=not (n and r and d and f and u)):
                            reg.upsert_chemo(rb.Chemotherapy(n, r, d, f, u))
                            bank.upsert_regimen(reg)
                            st.success("Added.")
                            _rerun()
                    with tab2:
                        cat, name_keys, display = agent_catalog_options(bank)
                        if not name_keys:
                            st.caption("No saved agents yet.")
                        else:
                            i_name = st.selectbox("Existing agent", options=range(len(display)), format_func=lambda i: display[i], key="cat_name")
                            variants = cat[name_keys[i_name]]
                            labels = [f"{v.route} | {v.dose} | {v.frequency} | {v.duration}" for v in variants]
                            i_var = st.selectbox("Variant", options=range(len(labels)), format_func=lambda i: labels[i], key="cat_var")
                            base = variants[i_var]
                            st.caption("You can tweak before saving.")
                            name = st.text_input("Agent name", value=base.name, key="c_name")
                            route = st.selectbox("Route", ["IV","PO","SQ","IM","IT"], index=["IV","PO","SQ","IM","IT"].index(base.route) if base.route in ["IV","PO","SC","IM","IT"] else 0, key="c_route")
                            dose = st.text_input("Dose", value=base.dose, key="c_dose")
                            freq = st.text_input("Frequency", value=base.frequency, key="c_freq")
                            dur  = st.text_input("Duration", value=base.duration, key="c_dur")
                            if st.button("Add from existing"):
                                reg.upsert_chemo(rb.Chemotherapy(name, route, dose, freq, dur))
                                bank.upsert_regimen(reg)
                                st.success("Added.")
                                _rerun()

    with colB:
        st.markdown("### Danger zone")
        names = get_regimens(bank)
        if names:
            sel_del = st.selectbox("Select regimen to delete", options=names, key="del_sel")
            if st.button("Delete regimen"):
                if bank.delete_regimen(sel_del):
                    st.error(f"Deleted '{sel_del}'.")
                    _rerun()

# ---------- Calendar View ----------
if view == "Calendar":
    st.header("Generate Calendar")

    sel = regimen_select_ui(bank, key="reg_for_cal")
    if not sel:
        st.stop()
    reg = bank.get_regimen(sel)

    with st.container(border=True):
        c1, c2, c3 = st.columns(3)
        with c1:
            start = st.date_input("Day 1 (can be in past)", value=dt.date.today())
        with c2:
            cycle_len = st.number_input("Cycle length (days)", min_value=1, step=1, value=28)
        with c3:
            cycle_label = cycle_label_ui()  # Cycle # or Induction

    if st.button("Create calendar", type="primary"):
        cal_txt = rb.make_calendar_text(reg, start, int(cycle_len), cycle_label)

        st.subheader("Preview")
        st.code(cal_txt, language="text")

        # Downloads
        safe_name = "".join(c if c.isalnum() or c in ("-", "_") else "_" for c in reg.name)
        base = f"{safe_name}_{cycle_label.replace(' ','').lower()}_{start.isoformat()}"

        st.download_button(
            label="Download TXT",
            data=cal_txt.encode("utf-8"),
            file_name=f"{base}.txt",
            mime="text/plain"
        )

        docx_bytes = export_docx_bytes(reg, start, int(cycle_len), cycle_label)
        if docx_bytes:
            st.download_button(
                label="Download DOCX",
                data=docx_bytes,
                file_name=f"{base}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.caption("Install python-docx to enable DOCX export: `pip install python-docx`")

    # Quick reference of selected regimen
    with st.expander("Regimen details"):
        st.write(f"**Name:** {reg.name}")
        if reg.disease_state:
            st.write(f"**Disease State:** {reg.disease_state}")
        if not reg.therapies:
            st.write("_No agents saved yet._")
        else:
            for t in reg.therapies:
                st.write(f"- **{t.name}** — {t.route} • {t.dose} • {t.frequency} • {t.duration}")
